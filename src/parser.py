import io
import PyPDF2
import docx

def extract_text_from_pdf(file) -> str:
    try:
        file.seek(0)
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    except Exception as e:
        raise ValueError(f"PDF 解析失败: {e}")

def extract_text_from_docx(file) -> str:
    try:
        file_bytes = file.getvalue()
        doc = docx.Document(io.BytesIO(file_bytes))

        # 1. 提取所有段落文本
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        # 2. 提取所有表格中的文本
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    tables_text.append(" | ".join(row_text))

        # 3. 合并
        all_text = "\n".join(paragraphs + tables_text)

        if not all_text.strip():
            raise ValueError("未提取到文本。请确认文档内容不是纯图片或扫描件，建议转为 PDF 上传。")
        return all_text.strip()
    except Exception as e:
        raise ValueError(f"DOCX 解析失败: {e}")

def parse_resume(uploaded_file) -> str:
    if uploaded_file is None:
        return ""
    file_type = uploaded_file.name.split(".")[-1].lower()
    if file_type == "pdf":
        return extract_text_from_pdf(uploaded_file)
    elif file_type == "docx":
        return extract_text_from_docx(uploaded_file)
    else:
        raise ValueError("不支持的文件类型，请上传 PDF 或 DOCX")
